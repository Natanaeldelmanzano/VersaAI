from typing import List, Optional, Dict, BinaryIO
from sqlalchemy.orm import Session
from ..models.knowledge_base import KnowledgeBase, Document, DocumentChunk, DocumentType, DocumentStatus
from ..core.config import settings
from .rag_service import rag_service
import os
import uuid
import time
import logging
from pathlib import Path
import aiofiles
import PyPDF2
import docx
import markdown
import json
import csv
from io import StringIO

logger = logging.getLogger(__name__)

class DocumentService:
    def __init__(self):
        self.upload_dir = Path(settings.UPLOAD_DIR)
        self.upload_dir.mkdir(exist_ok=True)
        self.max_file_size = settings.MAX_FILE_SIZE_MB * 1024 * 1024  # Convert MB to bytes
        self.chunk_size = 1000
        self.chunk_overlap = 200
    
    def get_supported_file_types(self) -> List[str]:
        """Get list of supported file types"""
        return ["txt", "pdf", "docx", "html", "md", "csv", "json"]
    
    def detect_document_type(self, filename: str) -> DocumentType:
        """Detect document type from filename"""
        extension = filename.lower().split('.')[-1]
        
        type_mapping = {
            'txt': DocumentType.TEXT,
            'pdf': DocumentType.PDF,
            'docx': DocumentType.DOCX,
            'html': DocumentType.HTML,
            'htm': DocumentType.HTML,
            'md': DocumentType.MARKDOWN,
            'markdown': DocumentType.MARKDOWN,
            'csv': DocumentType.CSV,
            'json': DocumentType.JSON
        }
        
        return type_mapping.get(extension, DocumentType.TEXT)
    
    def validate_file(self, filename: str, file_size: int) -> tuple[bool, str]:
        """Validate uploaded file"""
        # Check file size
        if file_size > self.max_file_size:
            return False, f"File size ({file_size} bytes) exceeds maximum allowed size ({self.max_file_size} bytes)"
        
        # Check file extension
        extension = filename.lower().split('.')[-1]
        if extension not in self.get_supported_file_types():
            return False, f"File type '{extension}' is not supported. Supported types: {', '.join(self.get_supported_file_types())}"
        
        return True, "File is valid"
    
    async def save_uploaded_file(self, file_content: bytes, original_filename: str) -> tuple[str, str]:
        """Save uploaded file to disk"""
        try:
            # Generate unique filename
            file_extension = original_filename.split('.')[-1]
            unique_filename = f"{uuid.uuid4()}.{file_extension}"
            file_path = self.upload_dir / unique_filename
            
            # Save file
            async with aiofiles.open(file_path, 'wb') as f:
                await f.write(file_content)
            
            logger.info(f"File saved: {unique_filename}")
            return str(file_path), unique_filename
            
        except Exception as e:
            logger.error(f"Error saving file: {e}")
            raise Exception(f"Failed to save file: {str(e)}")
    
    def extract_text_from_file(self, file_path: str, document_type: DocumentType) -> str:
        """Extract text content from file based on type"""
        try:
            if document_type == DocumentType.TEXT:
                with open(file_path, 'r', encoding='utf-8') as f:
                    return f.read()
            
            elif document_type == DocumentType.PDF:
                return self._extract_text_from_pdf(file_path)
            
            elif document_type == DocumentType.DOCX:
                return self._extract_text_from_docx(file_path)
            
            elif document_type == DocumentType.HTML:
                return self._extract_text_from_html(file_path)
            
            elif document_type == DocumentType.MARKDOWN:
                return self._extract_text_from_markdown(file_path)
            
            elif document_type == DocumentType.CSV:
                return self._extract_text_from_csv(file_path)
            
            elif document_type == DocumentType.JSON:
                return self._extract_text_from_json(file_path)
            
            else:
                # Fallback to text
                with open(file_path, 'r', encoding='utf-8') as f:
                    return f.read()
                    
        except Exception as e:
            logger.error(f"Error extracting text from {file_path}: {e}")
            raise Exception(f"Failed to extract text: {str(e)}")
    
    def _extract_text_from_pdf(self, file_path: str) -> str:
        """Extract text from PDF file"""
        text = ""
        with open(file_path, 'rb') as f:
            pdf_reader = PyPDF2.PdfReader(f)
            for page in pdf_reader.pages:
                text += page.extract_text() + "\n"
        return text
    
    def _extract_text_from_docx(self, file_path: str) -> str:
        """Extract text from DOCX file"""
        doc = docx.Document(file_path)
        text = ""
        for paragraph in doc.paragraphs:
            text += paragraph.text + "\n"
        return text
    
    def _extract_text_from_html(self, file_path: str) -> str:
        """Extract text from HTML file"""
        with open(file_path, 'r', encoding='utf-8') as f:
            html_content = f.read()
        
        # Simple HTML tag removal (for production, consider using BeautifulSoup)
        import re
        text = re.sub(r'<[^>]+>', '', html_content)
        text = re.sub(r'\s+', ' ', text)
        return text.strip()
    
    def _extract_text_from_markdown(self, file_path: str) -> str:
        """Extract text from Markdown file"""
        with open(file_path, 'r', encoding='utf-8') as f:
            md_content = f.read()
        
        # Convert markdown to HTML then extract text
        html = markdown.markdown(md_content)
        import re
        text = re.sub(r'<[^>]+>', '', html)
        text = re.sub(r'\s+', ' ', text)
        return text.strip()
    
    def _extract_text_from_csv(self, file_path: str) -> str:
        """Extract text from CSV file"""
        text_parts = []
        with open(file_path, 'r', encoding='utf-8') as f:
            csv_reader = csv.reader(f)
            headers = next(csv_reader, None)
            if headers:
                text_parts.append("Headers: " + ", ".join(headers))
            
            for i, row in enumerate(csv_reader):
                if i >= 100:  # Limit to first 100 rows
                    break
                text_parts.append("Row: " + ", ".join(row))
        
        return "\n".join(text_parts)
    
    def _extract_text_from_json(self, file_path: str) -> str:
        """Extract text from JSON file"""
        with open(file_path, 'r', encoding='utf-8') as f:
            data = json.load(f)
        
        # Convert JSON to readable text
        return json.dumps(data, indent=2, ensure_ascii=False)
    
    def split_text_into_chunks(self, text: str, chunk_size: int = None, chunk_overlap: int = None) -> List[str]:
        """Split text into chunks for processing"""
        chunk_size = chunk_size or self.chunk_size
        chunk_overlap = chunk_overlap or self.chunk_overlap
        
        if len(text) <= chunk_size:
            return [text]
        
        chunks = []
        start = 0
        
        while start < len(text):
            end = start + chunk_size
            
            # Try to break at sentence or word boundary
            if end < len(text):
                # Look for sentence boundary
                sentence_end = text.rfind('.', start, end)
                if sentence_end > start + chunk_size // 2:
                    end = sentence_end + 1
                else:
                    # Look for word boundary
                    word_end = text.rfind(' ', start, end)
                    if word_end > start + chunk_size // 2:
                        end = word_end
            
            chunk = text[start:end].strip()
            if chunk:
                chunks.append(chunk)
            
            # Move start position with overlap
            start = end - chunk_overlap
            if start >= len(text):
                break
        
        return chunks
    
    async def process_document(
        self, 
        db: Session, 
        document_id: int
    ) -> bool:
        """Process a document: extract text, create chunks, generate embeddings"""
        try:
            # Get document
            document = db.query(Document).filter(Document.id == document_id).first()
            if not document:
                logger.error(f"Document {document_id} not found")
                return False
            
            # Update status to processing
            document.status = DocumentStatus.PROCESSING
            db.commit()
            
            start_time = time.time()
            
            # Extract text content
            logger.info(f"Extracting text from document {document.filename}")
            text_content = self.extract_text_from_file(document.file_path, document.document_type)
            
            if not text_content.strip():
                document.status = DocumentStatus.ERROR
                document.error_message = "No text content extracted from file"
                db.commit()
                return False
            
            # Save extracted content
            document.content = text_content
            
            # Split into chunks
            logger.info(f"Splitting document into chunks")
            chunks = self.split_text_into_chunks(text_content)
            
            # Delete existing chunks
            db.query(DocumentChunk).filter(DocumentChunk.document_id == document_id).delete()
            
            # Create new chunks with embeddings
            logger.info(f"Creating {len(chunks)} chunks with embeddings")
            for i, chunk_text in enumerate(chunks):
                # Generate embedding
                embedding = rag_service.generate_embedding(chunk_text)
                
                # Create chunk
                chunk = DocumentChunk(
                    document_id=document_id,
                    chunk_index=i,
                    content=chunk_text,
                    embedding=embedding,
                    token_count=len(chunk_text.split())
                )
                
                db.add(chunk)
            
            # Update document status
            processing_time = time.time() - start_time
            document.status = DocumentStatus.PROCESSED
            document.chunks_count = len(chunks)
            document.processing_time_seconds = processing_time
            document.processed_at = datetime.utcnow()
            
            # Update knowledge base statistics
            kb = document.knowledge_base
            kb.total_chunks = db.query(DocumentChunk).join(Document).filter(
                Document.knowledge_base_id == kb.id
            ).count()
            kb.last_updated = datetime.utcnow()
            
            db.commit()
            
            logger.info(f"Document {document.filename} processed successfully in {processing_time:.2f}s")
            return True
            
        except Exception as e:
            logger.error(f"Error processing document {document_id}: {e}")
            
            # Update document status to error
            try:
                document = db.query(Document).filter(Document.id == document_id).first()
                if document:
                    document.status = DocumentStatus.ERROR
                    document.error_message = str(e)
                    db.commit()
            except:
                pass
            
            return False
    
    async def upload_and_process_document(
        self,
        db: Session,
        knowledge_base_id: int,
        file_content: bytes,
        original_filename: str,
        title: Optional[str] = None
    ) -> Optional[Document]:
        """Upload and process a document"""
        try:
            # Validate file
            is_valid, message = self.validate_file(original_filename, len(file_content))
            if not is_valid:
                raise Exception(message)
            
            # Save file
            file_path, filename = await self.save_uploaded_file(file_content, original_filename)
            
            # Detect document type
            document_type = self.detect_document_type(original_filename)
            
            # Create document record
            document = Document(
                filename=filename,
                original_filename=original_filename,
                title=title or original_filename,
                knowledge_base_id=knowledge_base_id,
                document_type=document_type,
                file_size=len(file_content),
                file_path=file_path,
                status=DocumentStatus.PENDING
            )
            
            db.add(document)
            db.commit()
            db.refresh(document)
            
            # Process document asynchronously
            success = await self.process_document(db, document.id)
            
            if success:
                logger.info(f"Document {original_filename} uploaded and processed successfully")
                return document
            else:
                logger.error(f"Failed to process document {original_filename}")
                return document  # Return even if processing failed, so user can see the error
                
        except Exception as e:
            logger.error(f"Error uploading document: {e}")
            raise Exception(f"Failed to upload document: {str(e)}")
    
    def delete_document(self, db: Session, document_id: int) -> bool:
        """Delete a document and its associated files"""
        try:
            document = db.query(Document).filter(Document.id == document_id).first()
            if not document:
                return False
            
            # Delete file from disk
            try:
                if os.path.exists(document.file_path):
                    os.remove(document.file_path)
            except Exception as e:
                logger.warning(f"Failed to delete file {document.file_path}: {e}")
            
            # Delete chunks
            db.query(DocumentChunk).filter(DocumentChunk.document_id == document_id).delete()
            
            # Delete document
            db.delete(document)
            db.commit()
            
            logger.info(f"Document {document.filename} deleted successfully")
            return True
            
        except Exception as e:
            logger.error(f"Error deleting document {document_id}: {e}")
            db.rollback()
            return False

# Global instance
document_service = DocumentService()